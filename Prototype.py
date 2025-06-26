import streamlit as st
import pandas as pd
import numpy as np
import scipy.stats as stats
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from PIL import Image
import os
import base64

# 🌈 ShadyTable 主題配色與 icon 定義
PRIMARY_COLOR = "#2E6F72"     # 深藍綠
SECONDARY_COLOR = "#5F6C7B"   # 明智灰
ACCENT_COLOR = "#B8F2E6"      # AI萌綠
HIGHLIGHT_COLOR = "#FFD23F"   # 檸檬黃
BG_COLOR = "#FAFAFA"          # 奶油白

ICON_UPLOAD = "📤"
ICON_CONFIG = "🔧"
ICON_STATS = "📊"
ICON_TABLE = "📄"
ICON_EXPORT = "💾"
ICON_WORD = "📝"
ICON_EXCEL = "📥"

# 頁面設定
st.set_page_config(page_title="ShadyTable", layout="wide", page_icon="🧠")

# Logo 與標題
def render_svg(svg_path):
    with open(svg_path, "r", encoding="utf-8") as f:
        svg_content = f.read()
    b64 = base64.b64encode(svg_content.encode("utf-8")).decode("utf-8")
    svg_html = f'<img src="data:image/svg+xml;base64,{b64}" style="width:100%;"/>'
    st.markdown(svg_html, unsafe_allow_html=True)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
svg_path = os.path.join(BASE_DIR, "assets", "shady_logo_banner.svg")

if os.path.exists(svg_path):
    render_svg(svg_path)
    st.markdown(
        f"""
        <div style='text-align: center;'>
            <h1 style='color:{PRIMARY_COLOR}; font-size: 36px;'>ShadyTable</h1>
            <h3 style='color:{SECONDARY_COLOR};'>「學長，我來」</h3>
        </div>
        """,
        unsafe_allow_html=True
    )
else:
    st.warning("⚠️ 無法載入 Logo，請確認圖片路徑正確。")

st.markdown(
    f"<h1 style='color:{PRIMARY_COLOR}'>🧠 ShadyTable - 智能 Table 1 無痛生成器</h1>",
    unsafe_allow_html=True
)

# Step 1: 上傳資料
st.markdown(f"### {ICON_UPLOAD} <span style='color:{SECONDARY_COLOR}'>Step 1: 上傳資料檔案</span>", unsafe_allow_html=True)
st.markdown("⚠️ **請務必移除所有個資欄位（如姓名、病歷號等），避免違反資料安全規範！**")
uploaded_file = st.file_uploader("請選擇 Excel 或 CSV 檔", type=["csv", "xlsx"])

if uploaded_file:
    try:
        raw_df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith(".csv") else pd.read_excel(uploaded_file)
        raw_df.replace(["NA", "na", "Na", ""], np.nan, inplace=True)

        df = raw_df.copy()

        personal_info_keywords = ["name", "Name", "姓名", "chart", "Chart_No", "病歷號"]
        alert_cols = [col for col in df.columns if any(key in col for key in personal_info_keywords)]
        if alert_cols:
            st.warning(f"⚠️ 偵測到可能含有個資的欄位：{', '.join(alert_cols)}，請確認是否已去識別化。")

        st.session_state["raw_df"] = raw_df  # 儲存未處理原始資料
        st.success("✅ 資料上傳成功，缺值已自動處理。")
    except Exception as e:
        st.error(f"❌ 檔案載入錯誤：{e}")
        df = None
else:
    df = None


# Step 2: 選擇變項
if df is not None:
    st.markdown(f"### {ICON_CONFIG} <span style='color:{SECONDARY_COLOR}'>Step 2: 選擇變項</span>", unsafe_allow_html=True)
    all_columns = df.columns.tolist()

    with st.form("var_select_form"):
        group_var = st.selectbox("🧩 分組變項（限兩組）", ["(No Grouping)"] + all_columns)
        cat_vars = st.multiselect("🔠 類別變項", all_columns)
        cont_vars = st.multiselect("🔢 連續變項", [col for col in all_columns if col not in cat_vars])
        fill_na = st.checkbox("🩹 自動填補缺值（平均/眾數）", value=True)
        submitted = st.form_submit_button("✅ 確認")

    if submitted:
        if fill_na:
            df = df.copy()
            for col in cont_vars:
                df[col].fillna(df[col].mean(), inplace=True)
            for col in cat_vars:
                df[col].fillna(df[col].mode().iloc[0], inplace=True)
            st.info("已填補缺值。")

        st.session_state["df"] = df
        st.session_state["group"] = None if group_var == "(No Grouping)" else group_var
        st.session_state["cat"] = cat_vars
        st.session_state["cont"] = cont_vars
        st.success("✅ 變項選擇完成。")

# Step 3: 統計分析與呈現
        st.markdown(f"### {ICON_STATS} <span style='color:{SECONDARY_COLOR}'>Step 3: 統計分析與結果呈現</span>", unsafe_allow_html=True)

    if "df" in st.session_state:
        df = st.session_state["df"]
        group_var = st.session_state["group"]
        cat_vars = st.session_state["cat"]
        cont_vars = st.session_state["cont"]

                # 計算每個 group 的有效樣本數
        group_labels = []
        if group_var:
            group_counts = df[group_var].value_counts(dropna=False).to_dict()
            group_labels = [f"{g} (n={group_counts[g]})" for g in group_counts.keys()]
        else:
            group_labels = ["Overall"]

        result_rows = []

        def normality_test(data):
            if len(data.dropna()) < 3:
                return False
            p = stats.shapiro(data.dropna())[1]
            return p > 0.05

        def format_continuous(series):
            mean = series.mean()
            std = series.std()
            median = series.median()
            q1 = series.quantile(0.25)
            q3 = series.quantile(0.75)
            return f"{mean:.1f} ± {std:.1f}", f"{median:.1f} [{q1:.1f}-{q3:.1f}]"

        def format_p(p):
            if p < 0.001:
                return "<0.001***"
            elif p < 0.01:
                return f"{p:.3f}**"
            elif p < 0.05:
                return f"{p:.3f}*"
            else:
                return f"{p:.3f}"

        # 類別變項
        for var in cat_vars:
            total = len(df)
            raw_df = st.session_state.get("raw_df", df)
            na_pct = raw_df[var].isna().mean() * 100
            row_header = {
                "Variable": f"**{var}**",
                "Missing (%)": f"{na_pct:.1f}%",
                "P": "",
                "Normality": "-",
                "Method": ""
            }
            sub_rows = []

            if group_var:
                ct = pd.crosstab(df[var], df[group_var])
                if ct.shape[1] == 2 and (ct.values < 5).sum() > 0:
                    _, p = stats.fisher_exact(ct.values)
                    row_header["Method"] = "Fisher's exact"
                else:
                    _, p, _, _ = stats.chi2_contingency(ct)
                    row_header["Method"] = "Chi-square"
                row_header["P"] = format_p(p)

                for level in ct.index:
                    sub_row = {"Variable": f"　{level}", "Missing (%)": "", "P": "", "Normality": "", "Method": ""}
                    for g in ct.columns:
                        count = ct.loc[level, g]
                        total = ct[g].sum()
                        label = f"{g} (n={group_counts[g]})"
                        sub_row[label] = f"{count} ({count/total*100:.1f}%)"
                    sub_rows.append(sub_row)
            else:
                counts = df[var].value_counts(dropna=False)
                for level, count in counts.items():
                    sub_rows.append({
                        "Variable": f"　{level}",
                        "Missing (%)": "", "P": "", "Normality": "", "Method": "",
                        "Overall": f"{count} ({count/len(df)*100:.1f}%)"
                    })
                row_header["P"] = "-"
                row_header["Method"] = "-"

            result_rows.append(row_header)
            result_rows.extend(sub_rows)

        # 連續變項
        for var in cont_vars:
            total = len(df)
            raw_df = st.session_state.get("raw_df", df)
            na_pct = raw_df[var].isna().mean() * 100
            s = df[var]
            is_normal = normality_test(s)
            method = "t-test" if is_normal else "Mann–Whitney U"

            row_header = {
                "Variable": f"**{var}**",
                "Missing (%)": f"{na_pct:.1f}%",
                "P": "",
                "Normality": "Yes" if is_normal else "No",
                "Method": method
            }
            row = {"Variable": "　Value", "Missing (%)": "", "P": "", "Normality": "", "Method": ""}

            if group_var:
                group_samples = [df[df[group_var] == g][var].dropna() for g in group_counts.keys()]
                for g, s_group in zip(group_counts.keys(), group_samples):
                    mean_std, med_iqr = format_continuous(s_group)
                    label = f"{g} (n={group_counts[g]})"
                    row[label] = mean_std if is_normal else med_iqr

                if len(group_samples) == 2:
                    if is_normal:
                        _, p = stats.ttest_ind(*group_samples, nan_policy="omit")
                    else:
                        _, p = stats.mannwhitneyu(*group_samples)
                    row_header["P"] = format_p(p)
                else:
                    row_header["P"] = "-"
            else:
                mean_std, med_iqr = format_continuous(s.dropna())
                row["Overall"] = mean_std if is_normal else med_iqr
                row_header["P"] = "-"

            result_rows.append(row_header)
            result_rows.append(row)

        table1 = pd.DataFrame(result_rows)

        # 重排欄位：Variable, [Group(s)], Missing (%), P, Normality, Method
        fixed_cols = ["Variable", "Missing (%)", "P", "Normality", "Method"]
        dynamic_cols = [col for col in table1.columns if col not in fixed_cols]
        col_order = ["Variable"] + dynamic_cols + ["Missing (%)", "P", "Normality", "Method"]
        table1 = table1[col_order]

        st.subheader("📄 Table 1 統計摘要")
        st.dataframe(table1, use_container_width=True)

        # Step 4: 匯出
        st.markdown("---")
        st.markdown(f"### {ICON_EXPORT} <span style='color:{SECONDARY_COLOR}'>Step 4: 匯出結果</span>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)

        with col1:
            def export_to_excel(df):
                output = BytesIO()
                df.to_excel(output, index=False)
                output.seek(0)
                return output

            excel_data = export_to_excel(table1)
            st.download_button("📥 匯出 Excel", data=excel_data, file_name="Table1_Shady.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        with col2:
            def export_to_word(df):
                doc = Document()
                doc.add_heading("Table 1 Summary", level=1)
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = col

                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)

                word_stream = BytesIO()
                doc.save(word_stream)
                word_stream.seek(0)
                return word_stream

            word_data = export_to_word(table1)
            st.download_button("📝 匯出 Word", data=word_data, file_name="Table1_Shady.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Step 5: 產出文字摘要
        st.markdown("---")
        st.markdown(f"### 🤖 <span style='color:{SECONDARY_COLOR}'>Step 5: 自動產出 Results 段落</span>", unsafe_allow_html=True)

        if st.button("🧠 使用 ChatGPT 撰寫結果摘要"):
            with st.spinner("🤖 正在請 ChatGPT 撰寫摘要，請稍候..."):
                try:
                    from openai import OpenAI
                    from dotenv import load_dotenv
                    load_dotenv(dotenv_path="./keys.env")
                    import os

                    client = OpenAI(api_key= st.secrets["OPENAI_API_KEY"])

                    prompt = f"""你是一位協助撰寫醫學論文的AI助手。請根據下列Table 1統計結果，撰寫一段英文Results摘要段落，不需要推論，只需要描述各變項的敘述統計與p值的差異即可。\n\n{table1.to_string(index=False)}\n\n請用正式英文寫一段200字內的段落作為Results。"""

                    response = client.chat.completions.create(
                        model="gpt-4",
                        messages=[
                            {"role": "system", "content": "You are a medical research assistant specialized in academic writing."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.3,
                        max_tokens=400
                    )

                    result_text = response.choices[0].message.content
                    st.markdown("✅ **自動產出 Results 段落如下：**")
                    st.text_area("Results", value=result_text, height=200)

                    def export_results_to_word(text):
                        doc = Document()
                        doc.add_heading("Results", level=1)
                        para = doc.add_paragraph(text)
                        para.style.font.size = Pt(12)
                        word_stream = BytesIO()
                        doc.save(word_stream)
                        word_stream.seek(0)
                        return word_stream

                    st.markdown("### 📝 下載 Results Word 檔")
                    word_result = export_results_to_word(result_text)
                    st.download_button(
                        label="📝 下載 Results.docx",
                        data=word_result,
                        file_name="Results_Shady.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"❌ 發生錯誤：{e}")


